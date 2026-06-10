"""Customer description extraction and Word document export."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from rvtools_processor import Workbook, discover_clusters
from vm_analytics import extract_vm_analytics

MAX_APPLICATIONS = 50


def _safe_str(value: Any) -> str:
    if pd.isna(value) or value is None:
        return ""
    return str(value).strip()


def _storage_summary(workbook: Workbook) -> dict[str, Any]:
    if "vDatastore" not in workbook:
        return {"types": [], "summary": "No storage information available"}

    df = workbook["vDatastore"]
    if "Type" not in df.columns:
        return {"types": [], "summary": "No storage information available"}

    types: list[dict[str, Any]] = []
    for storage_type, group in df.groupby(df["Type"].astype(str)):
        if storage_type in ("nan", ""):
            continue
        entry: dict[str, Any] = {
            "type": storage_type,
            "count": int(len(group)),
        }
        if "Capacity MiB" in group.columns:
            capacity = pd.to_numeric(group["Capacity MiB"], errors="coerce").sum()
            if pd.notna(capacity) and capacity > 0:
                entry["capacity_tib"] = round(capacity / 1024 / 1024, 1)
        types.append(entry)

    types.sort(key=lambda item: item["type"])
    parts = []
    for item in types:
        text = f"{item['type']} ({item['count']} Datastores"
        if "capacity_tib" in item:
            text += f", approx. {item['capacity_tib']} TiB"
        text += ")"
        parts.append(text)

    summary = ", ".join(parts) if parts else "No storage information available"
    return {"types": types, "summary": summary}


def _node_cores_analysis(workbook: Workbook) -> dict[str, Any]:
    if "vHost" not in workbook:
        return {
            "nodes": [],
            "total_cores": 0,
            "total_cpus": 0,
            "node_count": 0,
            "avg_cores": 0,
            "min_cores": 0,
            "max_cores": 0,
            "summary": "No host data available",
        }

    df = workbook["vHost"]
    nodes: list[dict[str, Any]] = []
    core_values: list[int] = []

    for _, row in df.iterrows():
        host = _safe_str(row.get("Host"))
        if not host:
            continue

        cores = pd.to_numeric(row.get("# Cores"), errors="coerce")
        cpus = pd.to_numeric(row.get("# CPU"), errors="coerce")
        cores_per_cpu = pd.to_numeric(row.get("Cores per CPU"), errors="coerce")
        memory = pd.to_numeric(row.get("# Memory"), errors="coerce")

        core_count = int(cores) if pd.notna(cores) else 0
        core_values.append(core_count)

        nodes.append({
            "host": host,
            "cluster": _safe_str(row.get("Cluster")),
            "cores": core_count,
            "cpus": int(cpus) if pd.notna(cpus) else 0,
            "cores_per_cpu": int(cores_per_cpu) if pd.notna(cores_per_cpu) else 0,
            "memory_gib": round(memory / 1024, 1) if pd.notna(memory) and memory > 0 else 0,
            "cpu_model": _safe_str(row.get("CPU Model")),
        })

    nodes.sort(key=lambda item: item["host"].lower())
    total_cores = sum(core_values)
    node_count = len(nodes)

    if node_count:
        summary = (
            f"{node_count} nodes, {total_cores} cores total "
            f"(avg {round(total_cores / node_count, 1)} cores/node, "
            f"min {min(core_values)}, max {max(core_values)})"
        )
    else:
        summary = "No host data available"

    return {
        "nodes": nodes,
        "total_cores": total_cores,
        "total_cpus": sum(n["cpus"] for n in nodes),
        "node_count": node_count,
        "avg_cores": round(total_cores / node_count, 1) if node_count else 0,
        "min_cores": min(core_values) if core_values else 0,
        "max_cores": max(core_values) if core_values else 0,
        "summary": summary,
    }


def extract_customer_summary(workbook: Workbook) -> dict[str, Any]:
    clusters = discover_clusters(workbook)
    cluster_count = len(clusters)
    node_count = len(workbook["vHost"]) if "vHost" in workbook else 0
    vm_count = len(workbook["vInfo"]) if "vInfo" in workbook else 0

    storage = _storage_summary(workbook)
    node_cores = _node_cores_analysis(workbook)
    analytics = extract_vm_analytics(workbook)
    listed_apps = analytics["application_names"][:MAX_APPLICATIONS]
    total_apps = analytics["application_count"]

    cluster_names = [c["name"] for c in clusters]

    return {
        "cluster_count": cluster_count,
        "cluster_names": cluster_names,
        "node_count": node_count,
        "vm_count": vm_count,
        "storage": storage,
        "applications": listed_apps,
        "application_count": total_apps,
        "applications_truncated": total_apps > MAX_APPLICATIONS,
        "os_summary": analytics["os_summary"],
        "os_families": analytics["os_families"],
        "powered_on_count": analytics["powered_on_count"],
        "node_cores": node_cores,
    }


def generate_customer_doc(
    customer_name: str,
    address: str,
    summary: dict[str, Any],
    path: str | Path,
    *,
    contact_person: str = "",
    additional_info: str = "",
    software_info: str = "",
) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Customer Description", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Created on: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph("")

    doc.add_heading("Customer Details", level=1)
    doc.add_paragraph(f"Customer: {customer_name}")
    for line in address.splitlines():
        line = line.strip()
        if line:
            doc.add_paragraph(line)
    if contact_person:
        doc.add_paragraph(f"Contact person: {contact_person}")

    doc.add_paragraph("")

    doc.add_heading("Infrastructure Overview", level=1)

    node_cores = summary.get("node_cores", {})
    infra_table = doc.add_table(rows=5, cols=2)
    infra_table.style = "Table Grid"
    rows = [
        ("Number of clusters", str(summary["cluster_count"])),
        ("Number of nodes (ESXi hosts)", str(summary["node_count"])),
        ("Total CPU cores", str(node_cores.get("total_cores", 0))),
        ("Total VMs", str(summary["vm_count"])),
        ("Storage", summary["storage"]["summary"]),
    ]
    for idx, (label, value) in enumerate(rows):
        infra_table.rows[idx].cells[0].text = label
        infra_table.rows[idx].cells[1].text = value

    if summary["cluster_names"]:
        doc.add_paragraph("")
        doc.add_heading("Cluster", level=2)
        for name in summary["cluster_names"]:
            doc.add_paragraph(name, style="List Bullet")

    if node_cores.get("nodes"):
        doc.add_paragraph("")
        doc.add_heading("Nodes and CPU Cores", level=2)
        doc.add_paragraph(node_cores.get("summary", ""))
        node_table = doc.add_table(rows=len(node_cores["nodes"]) + 1, cols=4)
        node_table.style = "Table Grid"
        headers = ("Host", "Cluster", "Cores", "RAM (GiB)")
        for idx, header in enumerate(headers):
            node_table.rows[0].cells[idx].text = header
        for row_idx, node in enumerate(node_cores["nodes"], start=1):
            node_table.rows[row_idx].cells[0].text = node["host"]
            node_table.rows[row_idx].cells[1].text = node["cluster"] or "—"
            node_table.rows[row_idx].cells[2].text = str(node["cores"])
            node_table.rows[row_idx].cells[3].text = str(node["memory_gib"]) if node["memory_gib"] else "—"

    doc.add_paragraph("")
    doc.add_heading("Operating Systems (powered-on VMs)", level=1)
    doc.add_paragraph(
        f"Powered-on VMs: {summary.get('powered_on_count', summary['application_count'])}"
    )

    if summary.get("os_families"):
        doc.add_paragraph("")
        doc.add_heading("OS Families", level=2)
        for item in summary["os_families"]:
            doc.add_paragraph(f"{item['family']}: {item['count']} VMs", style="List Bullet")

    if summary.get("os_summary"):
        doc.add_paragraph("")
        doc.add_heading("Operating Systems (detail)", level=2)
        os_table = doc.add_table(rows=len(summary["os_summary"]) + 1, cols=2)
        os_table.style = "Table Grid"
        os_table.rows[0].cells[0].text = "Operating system"
        os_table.rows[0].cells[1].text = "Number of VMs"
        for idx, item in enumerate(summary["os_summary"], start=1):
            os_table.rows[idx].cells[0].text = item["os"]
            os_table.rows[idx].cells[1].text = str(item["count"])

    doc.add_paragraph("")
    doc.add_heading("Running Applications / VMs", level=1)
    app_count = summary["application_count"]
    doc.add_paragraph(
        f"{app_count} powered-on virtual machines were detected."
    )

    if summary["applications"]:
        for app in summary["applications"]:
            doc.add_paragraph(app, style="List Bullet")
        if summary["applications_truncated"]:
            remaining = app_count - len(summary["applications"])
            doc.add_paragraph(f"… and {remaining} more")
    else:
        doc.add_paragraph("No powered-on VMs found.")

    if software_info.strip():
        doc.add_paragraph("")
        doc.add_heading("Software Infrastructure", level=1)
        for line in software_info.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")

    if additional_info.strip():
        doc.add_paragraph("")
        doc.add_heading("Additional Information", level=1)
        doc.add_paragraph(additional_info.strip())

    path = Path(path)
    doc.save(path)
